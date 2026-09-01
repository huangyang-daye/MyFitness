"""知识库文档解析测试。"""

from io import BytesIO

import pytest

from myfitness.api.web import parse_multipart_file
from myfitness.rag.document_parser import DocumentParseError, parse_document
from myfitness.rag.knowledge_service import MAX_CONTENT_LEN


def test_parse_markdown():
    parsed = parse_document("减脂原则.md", "# 原则\n\n蛋白质每公斤 1.6g".encode())
    assert parsed.format == "md"
    assert parsed.title == "减脂原则"
    assert "1.6g" in parsed.content
    assert not parsed.truncated


def test_parse_gb18030_txt():
    parsed = parse_document("note.txt", "少油少盐多蔬菜".encode("gb18030"))
    assert parsed.format == "txt"
    assert "少油少盐" in parsed.content


def test_parse_html_strips_script():
    html = "<html><body><h1>饮食</h1><p>高蛋白</p><script>alert(1)</script></body></html>".encode()
    parsed = parse_document("tips.html", html)
    assert parsed.format == "html"
    assert "高蛋白" in parsed.content
    assert "alert" not in parsed.content
    assert "# 饮食" in parsed.content


def test_parse_rtf():
    parsed = parse_document("note.rtf", rb"{\rtf1\ansi Protein target 1.6g per kg}")
    assert parsed.format == "rtf"
    assert "Protein target" in parsed.content


def test_parse_docx():
    pytest.importorskip("docx")
    from docx import Document

    document = Document()
    document.add_heading("训练原则", level=1)
    document.add_paragraph("卧推时肩胛骨收紧。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "动作"
    table.cell(0, 1).text = "组数"
    table.cell(1, 0).text = "深蹲"
    table.cell(1, 1).text = "4"
    buffer = BytesIO()
    document.save(buffer)
    parsed = parse_document("训练计划.docx", buffer.getvalue())
    assert parsed.format == "docx"
    assert parsed.title == "训练计划"
    assert "肩胛骨" in parsed.content
    assert "深蹲" in parsed.content


def test_parse_pdf(monkeypatch):
    class FakePage:
        def extract_text(self) -> str:
            return "蛋白质每公斤体重 1.6g"

    class FakeReader:
        def __init__(self, _stream) -> None:
            self.pages = [FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    parsed = parse_document("plan.pdf", b"%PDF-1.4 fake-body")
    assert parsed.format == "pdf"
    assert "1.6g" in parsed.content
    assert "第 1 页" in parsed.content


def test_parse_legacy_doc_utf16():
    body = "减脂期蛋白质要达到每公斤 1.6 克以上。".encode("utf-16le")
    data = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 80 + body + b"\x00\x00"
    parsed = parse_document("原则.doc", data)
    assert parsed.format == "doc"
    assert "蛋白质" in parsed.content


def test_parse_rejects_empty():
    with pytest.raises(DocumentParseError, match="空"):
        parse_document("empty.md", b"")


def test_parse_rejects_unsupported():
    with pytest.raises(DocumentParseError, match="不支持"):
        parse_document("photo.png", b"\x89PNG\r\n\x1a\n" + b"\x00" * 20)


def test_parse_truncates_long_text():
    payload = ("原则\n" + ("蛋白质 " * 20_000)).encode("utf-8")
    parsed = parse_document("long.md", payload)
    assert parsed.truncated
    assert len(parsed.content) <= MAX_CONTENT_LEN


def test_parse_multipart_file_extracts_payload():
    boundary = "----TestBoundary"
    filename = "原则.md"
    content = "# hello\n"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        "Content-Type: text/markdown\r\n"
        "\r\n"
        f"{content}\r\n"
        f"--{boundary}--\r\n"
    ).encode()
    name, data = parse_multipart_file(body, f"multipart/form-data; boundary={boundary}")
    assert name == filename
    assert data == content.encode("utf-8")


def test_parse_multipart_requires_file_field():
    boundary = "----TestBoundary"
    body = (
        f"--{boundary}\r\n"
        'Content-Disposition: form-data; name="title"\r\n'
        "\r\n"
        "only-title\r\n"
        f"--{boundary}--\r\n"
    ).encode()
    with pytest.raises(ValueError, match="未找到"):
        parse_multipart_file(body, f"multipart/form-data; boundary={boundary}")
