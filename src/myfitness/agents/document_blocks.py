"""结构化文档块 — 用于原生 DOCX 写入与跨格式导出。"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

_BLOCK_TYPES = frozenset(
    {"title", "heading", "paragraph", "bullet_list", "numbered_list", "table"}
)


def parse_document_blocks(content: str, *, prefer_json: bool = False) -> list[dict[str, Any]]:
    """将 LLM JSON 或 Markdown 正文解析为统一块列表。"""
    text = str(content or "").strip()
    if not text:
        return []
    if prefer_json:
        blocks = _parse_json_blocks(text)
        if blocks:
            return blocks
    return markdown_to_blocks(text)


def blocks_to_markdown(blocks: list[dict[str, Any]]) -> str:
    lines: list[str] = []
    for block in blocks:
        btype = str(block.get("type") or block.get("kind") or "paragraph")
        if btype == "title":
            lines.extend(["# " + str(block.get("text", "")), ""])
        elif btype == "heading":
            level = max(1, min(int(block.get("level", 2)), 6))
            lines.extend(["#" * level + " " + str(block.get("text", "")), ""])
        elif btype == "paragraph":
            lines.extend([str(block.get("text", "")), ""])
        elif btype == "bullet_list":
            for item in block.get("items") or []:
                lines.append(f"- {item}")
            lines.append("")
        elif btype == "numbered_list":
            for idx, item in enumerate(block.get("items") or [], start=1):
                lines.append(f"{idx}. {item}")
            lines.append("")
        elif btype == "table":
            rows = block.get("rows") or []
            if not rows:
                continue
            col_count = max(len(row) for row in rows)
            header = rows[0]
            lines.append(
                "| "
                + " | ".join(str(header[i] if i < len(header) else "") for i in range(col_count))
                + " |"
            )
            lines.append("| " + " | ".join("---" for _ in range(col_count)) + " |")
            for row in rows[1:]:
                lines.append(
                    "| "
                    + " | ".join(str(row[i] if i < len(row) else "") for i in range(col_count))
                    + " |"
                )
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def write_docx_blocks(path: Path, blocks: list[dict[str, Any]]) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        from myfitness.agents.tools.document_tools import DocumentToolError

        raise DocumentToolError("写入 Word 需要安装 python-docx") from exc

    document = Document()
    for block in blocks:
        btype = str(block.get("type") or block.get("kind") or "paragraph")
        if btype == "title":
            document.add_heading(str(block.get("text", "")), level=0)
        elif btype == "heading":
            level = max(1, min(int(block.get("level", 2)), 3))
            document.add_heading(str(block.get("text", "")), level=level)
        elif btype == "paragraph":
            paragraph = document.add_paragraph(str(block.get("text", "")))
            for run in paragraph.runs:
                run.font.size = Pt(11)
        elif btype == "bullet_list":
            for item in block.get("items") or []:
                document.add_paragraph(str(item), style="List Bullet")
        elif btype == "numbered_list":
            for item in block.get("items") or []:
                document.add_paragraph(str(item), style="List Number")
        elif btype == "table":
            rows = block.get("rows") or []
            if not rows:
                continue
            col_count = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(col_count):
                    cell_text = str(row[c_idx]) if c_idx < len(row) else ""
                    table.cell(r_idx, c_idx).text = cell_text
    document.save(str(path))


def docx_to_preview_html(data: bytes) -> str:
    """将 DOCX 转为 HTML，供前端近似还原 Word 版式。"""
    try:
        import io

        import mammoth
    except ImportError:
        from myfitness.rag.document_parser import parse_document

        parsed = parse_document("preview.docx", data)
        escaped = (
            parsed.content.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br>")
        )
        return f'<div class="docx-plain-fallback"><p>{escaped}</p></div>'

    result = mammoth.convert_to_html(io.BytesIO(data))
    html = result.value.strip()
    if not html:
        return "<p>（文档为空）</p>"
    return html


def _parse_json_blocks(text: str) -> list[dict[str, Any]] | None:
    payload = _extract_json_object(text)
    if not isinstance(payload, dict):
        return None
    raw_blocks = payload.get("blocks")
    if not isinstance(raw_blocks, list):
        return None
    blocks: list[dict[str, Any]] = []
    for item in raw_blocks:
        if not isinstance(item, dict):
            continue
        btype = str(item.get("type") or item.get("kind") or "").strip().lower()
        if btype not in _BLOCK_TYPES:
            continue
        block: dict[str, Any] = {"type": btype}
        if btype in {"title", "heading", "paragraph"}:
            block["text"] = str(item.get("text", "")).strip()
            if not block["text"]:
                continue
            if btype == "heading":
                block["level"] = int(item.get("level", 2) or 2)
        elif btype in {"bullet_list", "numbered_list"}:
            items = item.get("items") or []
            if not isinstance(items, list):
                continue
            block["items"] = [str(x).strip() for x in items if str(x).strip()]
            if not block["items"]:
                continue
        elif btype == "table":
            rows = item.get("rows") or []
            if not isinstance(rows, list) or not rows:
                continue
            block["rows"] = [
                [str(cell).strip() for cell in row]
                for row in rows
                if isinstance(row, list) and any(str(cell).strip() for cell in row)
            ]
            if not block["rows"]:
                continue
        blocks.append(block)
    return blocks or None


def _extract_json_object(text: str) -> dict[str, Any] | None:
    raw = text.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    start = raw.find("{")
    end = raw.rfind("}")
    if start < 0 or end <= start:
        return None
    try:
        parsed = json.loads(raw[start : end + 1])
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


def write_pdf_blocks(path: Path, blocks: list[dict[str, Any]]) -> None:
    try:
        from fpdf import FPDF
        from fpdf.enums import XPos, YPos
    except ImportError as exc:
        from myfitness.agents.tools.document_tools import DocumentToolError

        raise DocumentToolError("写入 PDF 需要安装 fpdf2（pip install fpdf2）") from exc

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    font_name = _register_pdf_font(pdf)
    width = pdf.epw

    def writeln(text: str, *, size: int = 11, line_h: float = 7, gap: float = 2) -> None:
        cleaned = strip_inline_markdown(text)
        if not cleaned:
            return
        pdf.set_font(font_name, size=size)
        pdf.multi_cell(width, line_h, cleaned, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(gap)

    for block in blocks:
        kind = str(block.get("type") or block.get("kind") or "paragraph")
        if kind == "title":
            pdf.ln(1)
            writeln(str(block.get("text", "")), size=20, line_h=10, gap=5)
            continue
        if kind == "heading":
            level = max(1, min(int(block.get("level", 2) or 2), 3))
            sizes = {1: 16, 2: 14, 3: 12}
            writeln(str(block.get("text", "")), size=sizes[level], line_h=8, gap=4)
            continue
        if kind == "table":
            _write_pdf_table(pdf, font_name, width, block.get("rows") or [])
            pdf.ln(3)
            continue
        if kind == "bullet_list":
            for item in block.get("items") or []:
                cleaned = strip_inline_markdown(str(item))
                if not cleaned:
                    continue
                pdf.set_font(font_name, size=11)
                pdf.set_x(pdf.l_margin + 2)
                pdf.multi_cell(
                    width - 4,
                    6.5,
                    f"• {cleaned}",
                    new_x=XPos.LMARGIN,
                    new_y=YPos.NEXT,
                )
            pdf.ln(2)
            continue
        if kind == "numbered_list":
            for idx, item in enumerate(block.get("items") or [], start=1):
                cleaned = strip_inline_markdown(str(item))
                if not cleaned:
                    continue
                pdf.set_font(font_name, size=11)
                pdf.set_x(pdf.l_margin + 2)
                pdf.multi_cell(
                    width - 4,
                    6.5,
                    f"{idx}. {cleaned}",
                    new_x=XPos.LMARGIN,
                    new_y=YPos.NEXT,
                )
            pdf.ln(2)
            continue

        writeln(str(block.get("text", "")), size=11, line_h=7, gap=3)

    pdf.output(str(path))


def strip_inline_markdown(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"\*\*(.+?)\*\*", r"\1", value)
    value = re.sub(r"__(.+?)__", r"\1", value)
    value = re.sub(r"\*(.+?)\*", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
    return value.replace("**", "").replace("__", "").strip()


def _write_pdf_table(pdf: object, font_name: str, width: float, rows: list[list[Any]]) -> None:
    if not rows:
        return
    from fpdf.enums import TableBordersLayout, TableCellFillMode

    col_count = max(len(row) for row in rows)
    col_width = width / max(col_count, 1)
    normalized = [
        [strip_inline_markdown(str(row[c_idx]) if c_idx < len(row) else "") for c_idx in range(col_count)]
        for row in rows
    ]
    pdf.set_font(font_name, size=10)  # type: ignore[attr-defined]
    with pdf.table(  # type: ignore[attr-defined]
        width=width,
        col_widths=(col_width,) * col_count,
        line_height=7,
        borders_layout=TableBordersLayout.HORIZONTAL_LINES,
        cell_fill_mode=TableCellFillMode.ROWS,
        first_row_as_headings=False,
        text_align="LEFT",
        padding=2,
    ) as table:
        for row in normalized:
            pdf_row = table.row()
            for cell in row:
                pdf_row.cell(cell or " ")


def _register_pdf_font(pdf: object) -> str:
    font_path = _find_cjk_font()
    if not font_path:
        from myfitness.agents.tools.document_tools import DocumentToolError

        raise DocumentToolError(
            "写入 PDF 需要系统中文字体（Windows 请确认存在 simhei.ttf 或 msyh.ttc）"
        )
    try:
        pdf.add_font("MyFitnessCJK", "", font_path)  # type: ignore[attr-defined]
    except Exception as exc:  # noqa: BLE001
        from myfitness.agents.tools.document_tools import DocumentToolError

        raise DocumentToolError(f"加载 PDF 字体失败：{font_path}") from exc
    return "MyFitnessCJK"


def _find_cjk_font() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/msyhbd.ttc"),
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for candidate in candidates:
        if candidate.is_file():
            return str(candidate)
    return None


def markdown_to_blocks(content: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    lines = content.replace("\r\n", "\n").replace("\r", "\n").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.lstrip().startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            if level == 1 and not blocks:
                blocks.append({"type": "title", "text": text})
            else:
                blocks.append(
                    {
                        "type": "heading",
                        "level": max(1, min(level, 3)),
                        "text": text,
                    }
                )
            index += 1
            continue
        if re.match(r"^[-*+]\s+", line.strip()):
            items: list[str] = []
            while index < len(lines) and re.match(r"^[-*+]\s+", lines[index].strip()):
                items.append(re.sub(r"^[-*+]\s+", "", lines[index].strip()))
                index += 1
            blocks.append({"type": "bullet_list", "items": items})
            continue
        if re.match(r"^\d+\.\s+", line.strip()):
            items = []
            while index < len(lines) and re.match(r"^\d+\.\s+", lines[index].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[index].strip()))
                index += 1
            blocks.append({"type": "numbered_list", "items": items})
            continue
        if "|" in line and index + 1 < len(lines) and re.match(r"^\s*\|?[\s:-]+\|", lines[index + 1]):
            rows: list[list[str]] = []
            while index < len(lines) and "|" in lines[index]:
                if re.match(r"^\s*\|?[\s:-]+\|", lines[index]):
                    index += 1
                    continue
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                rows.append(cells)
                index += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines) and lines[index].strip() and not lines[index].lstrip().startswith("#"):
            if "|" in lines[index] or re.match(r"^[-*+]\s+", lines[index].strip()):
                break
            if re.match(r"^\d+\.\s+", lines[index].strip()):
                break
            paragraph_lines.append(lines[index].strip())
            index += 1
        blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines)})
    return blocks
